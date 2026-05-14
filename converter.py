"""Markdown 到 Word 文档的核心转换引擎。"""

import os
import re
import sys
import tempfile

import mistune
import requests
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

from config import FONT_CONFIG
from styles import (
    configure_page, configure_styles, set_run_font,
    apply_body_style, configure_table,
    get_heading_config, get_body_config,
)

_CN_DIGITS = '一二三四五六七八九十'


def _cn_number(n):
    """将整数转为中文数字（支持 1-99）。"""
    if n <= 10:
        return _CN_DIGITS[n - 1]
    tens, ones = divmod(n, 10)
    result = ''
    if tens > 1:
        result += _CN_DIGITS[tens - 1]
    result += '十'
    if ones:
        result += _CN_DIGITS[ones - 1]
    return result


class MarkdownToDocxConverter:
    """将 Markdown AST 转换为 python-docx 文档。"""

    # 标题编号格式（h1 无编号，h2-h5 按格式文档）
    HEADING_NUMBER_FORMATS = {
        2: lambda n: f'{_cn_number(n)}、',
        3: lambda n: f'（{_cn_number(n)}）',
        4: lambda n: f'{n}.',
        5: lambda n: f'（{n}）',
    }

    # 匹配任意级别的标题序号（用于检测已有序号，避免重复添加）
    _ANY_NUMBER_PATTERN = re.compile(
        r'^('
        r'[一二三四五六七八九十]+、'
        r'|（[一二三四五六七八九十]+）'
        r'|\d+[\.．]'
        r'|（\d+）'
        r')'
    )

    def __init__(self, document, image_base_path=None):
        self.document = document
        self.image_base_path = image_base_path or '.'
        self._current_paragraph = None
        # 标题计数器：每级独立计数，遇到上级标题时下级重置
        self._heading_counters = {2: 0, 3: 0, 4: 0, 5: 0}

    def convert(self, markdown_text):
        """解析 Markdown 文本并转换为 Word 文档元素。"""
        md = mistune.create_markdown(renderer='ast', plugins=['table', 'strikethrough'])
        tokens = md(markdown_text)
        self._process_tokens(tokens)

    def _process_tokens(self, tokens):
        """处理 AST token 列表。"""
        for token in tokens:
            self._process_token(token)

    def _process_token(self, token):
        """处理单个 AST token。"""
        t = token['type']

        if t == 'heading':
            self._handle_heading(token)
        elif t == 'paragraph':
            self._handle_paragraph(token)
        elif t == 'table':
            self._handle_table(token)
        elif t == 'list':
            self._handle_list(token)
        elif t == 'block_code':
            self._handle_block_code(token)
        elif t == 'block_quote':
            self._handle_block_quote(token)
        elif t == 'thematic_break':
            self._handle_thematic_break()
        elif t == 'blank_line':
            pass

    # ---- 块级处理 ----

    def _handle_heading(self, token):
        level = token['attrs']['level']
        cfg = get_heading_config(level)

        # 提取标题文本，检测是否已自带任何级别的序号
        heading_text = self._extract_text_from_children(token)
        already_numbered = bool(self._ANY_NUMBER_PATTERN.match(heading_text))

        # 更新标题计数器
        if level in self._heading_counters:
            self._heading_counters[level] += 1
        for lower in range(level + 1, 6):
            if lower in self._heading_counters:
                self._heading_counters[lower] = 0

        # 仅在标题未自带序号时添加自动编号
        prefix = ''
        if not already_numbered:
            counter = self._heading_counters.get(level, 0)
            fmt = self.HEADING_NUMBER_FORMATS.get(level)
            if fmt and counter:
                prefix = fmt(counter)

        p = self.document.add_paragraph()
        # h1 居中（文档标题），其他标题设置 28pt 行距、首行缩进 2 字符
        pf = p.paragraph_format
        if level == 1:
            pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pf.first_line_indent = Pt(0)
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        else:
            pf.line_spacing = Pt(cfg.get('line_spacing', 28))
            pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            pf.first_line_indent = Pt(cfg.get('first_line_indent', 32))

        # 添加编号前缀
        if prefix:
            run = p.add_run(prefix)
            set_run_font(run, cfg['name'], cfg['size'])

        self._current_paragraph = p
        self._process_inline_children(token.get('children', []))
        for run in p.runs:
            set_run_font(run, cfg['name'], cfg['size'])
        self._current_paragraph = None

    def _handle_paragraph(self, token):
        body_cfg = get_body_config()
        p = self.document.add_paragraph()
        apply_body_style(p)
        self._current_paragraph = p
        self._process_inline_children(token.get('children', []))
        # 设置字体
        for run in p.runs:
            set_run_font(run, body_cfg['name'], body_cfg['size'])
        self._current_paragraph = None

    def _handle_block_code(self, token):
        raw = token.get('raw', '')
        for line in raw.split('\n'):
            p = self.document.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)

    def _handle_block_quote(self, token):
        for child in token.get('children', []):
            if child['type'] == 'paragraph':
                p = self.document.add_paragraph()
                apply_body_style(p)
                p.paragraph_format.left_indent = Pt(32)
                self._current_paragraph = p
                self._process_inline_children(child.get('children', []))
                body_cfg = get_body_config()
                for run in p.runs:
                    set_run_font(run, body_cfg['name'], body_cfg['size'])
                self._current_paragraph = None

    def _handle_thematic_break(self):
        p = self.document.add_paragraph()
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('—' * 30)
        set_run_font(run, get_body_config()['name'], get_body_config()['size'])

    # ---- 列表处理 ----

    def _handle_list(self, token):
        ordered = token['attrs'].get('ordered', False)
        start = token['attrs'].get('start', 1)
        idx = start
        for i, child in enumerate(token.get('children', [])):
            if child['type'] == 'list_item':
                self._handle_list_item(child, ordered, idx)
                idx += 1

    def _handle_list_item(self, token, ordered, index):
        body_cfg = get_body_config()

        if ordered:
            bullet = f'{index}. '
        else:
            bullet = '● '

        p = self.document.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(body_cfg['first_line_indent'])
        p.paragraph_format.line_spacing = Pt(body_cfg['line_spacing'])
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY

        run = p.add_run(bullet)
        set_run_font(run, body_cfg['name'], body_cfg['size'])

        self._current_paragraph = p
        # list_item 的子节点可能是 block_text/paragraph（含行内子节点）
        for child in token.get('children', []):
            ct = child['type']
            if ct in ('block_text', 'paragraph'):
                self._process_inline_children(child.get('children', []))
            elif ct == 'list':
                self._handle_list(child)
        for r in p.runs[1:]:
            set_run_font(r, body_cfg['name'], body_cfg['size'])
        self._current_paragraph = None

    # ---- 表格处理 ----

    def _handle_table(self, token):
        rows_data = []
        for child in token.get('children', []):
            if child['type'] == 'table_head':
                for row_token in child.get('children', []):
                    if row_token['type'] == 'table_row':
                        row = self._extract_table_row(row_token, is_head=True)
                        rows_data.append(row)
            elif child['type'] == 'table_body':
                for row_token in child.get('children', []):
                    if row_token['type'] == 'table_row':
                        row = self._extract_table_row(row_token, is_head=False)
                        rows_data.append(row)

        if not rows_data:
            return

        num_rows = len(rows_data)
        num_cols = max(len(r) for r in rows_data)

        table = self.document.add_table(rows=num_rows, cols=num_cols)
        for i, row_data in enumerate(rows_data):
            for j, cell_data in enumerate(row_data):
                if j < num_cols:
                    cell = table.cell(i, j)
                    cell.text = ''
                    p = cell.paragraphs[0]
                    p.paragraph_format.first_line_indent = Pt(0)
                    body_cfg = get_body_config()
                    run = p.add_run(cell_data['text'])
                    set_run_font(run, body_cfg['name'], body_cfg['size'])
                    if cell_data['head']:
                        run.bold = True
        configure_table(table)

    def _extract_table_row(self, token, is_head=False):
        row = []
        for child in token.get('children', []):
            if child['type'] == 'table_cell':
                text = self._extract_text_from_children(child)
                row.append({'text': text, 'head': is_head})
        return row

    def _extract_text_from_children(self, token):
        """递归提取 token 子节点中的纯文本。"""
        if 'raw' in token:
            return token['raw']
        if 'children' not in token:
            return ''
        parts = []
        for child in token['children']:
            parts.append(self._extract_text_from_children(child))
        return ''.join(parts)

    # ---- 行内内容处理 ----

    def _process_inline_children(self, children):
        """处理行内 token 子节点列表。"""
        if not children:
            return
        for child in children:
            self._process_inline(child)

    def _process_inline(self, token):
        """处理单个行内 token。"""
        t = token['type']

        if t == 'text':
            raw = token.get('raw', '')
            if raw and self._current_paragraph:
                self._current_paragraph.add_run(raw)
        elif t == 'strong':
            self._add_formatted_run(token, bold=True)
        elif t == 'emphasis':
            self._add_formatted_run(token, italic=True)
        elif t == 'codespan':
            raw = token.get('raw', '')
            if raw and self._current_paragraph:
                run = self._current_paragraph.add_run(raw)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
        elif t == 'link':
            self._handle_link(token)
        elif t == 'image':
            self._handle_image(token)
        elif t == 'softbreak':
            if self._current_paragraph:
                self._current_paragraph.add_run(' ')
        elif t == 'linebreak':
            if self._current_paragraph:
                run = self._current_paragraph.add_run()
                run.add_break()
        elif t == 'inline_html':
            raw = token.get('raw', '')
            clean = re.sub(r'<[^>]+>', '', raw)
            if clean and self._current_paragraph:
                self._current_paragraph.add_run(clean)
        elif t == 'strikethrough':
            self._add_formatted_run(token, strike=True)

    def _add_formatted_run(self, token, bold=False, italic=False, strike=False):
        """添加带格式的 run（粗体/斜体/删除线）。"""
        text = self._extract_text_from_children(token)
        if text and self._current_paragraph:
            run = self._current_paragraph.add_run(text)
            if bold:
                run.bold = True
            if italic:
                run.italic = True
            if strike:
                run.font.strike = True

    def _handle_link(self, token):
        text = self._extract_text_from_children(token)
        if text and self._current_paragraph:
            run = self._current_paragraph.add_run(text)
            run.underline = True

    def _handle_image(self, token):
        attrs = token.get('attrs', {})
        url = attrs.get('url', '')
        alt = attrs.get('alt', '')

        is_url = url.startswith(('http://', 'https://'))
        img_path = self._resolve_image_path(url)
        if img_path and os.path.isfile(img_path):
            try:
                max_width = Inches(6)
                p = self.document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                run = p.add_run()
                run.add_picture(img_path, width=max_width)
                if alt:
                    caption_p = self.document.add_paragraph()
                    caption_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_p.paragraph_format.first_line_indent = Pt(0)
                    caption_run = caption_p.add_run(alt)
                    set_run_font(caption_run, get_body_config()['name'], Pt(12))
            finally:
                if is_url:
                    try:
                        os.remove(img_path)
                    except OSError:
                        pass

    def _resolve_image_path(self, url):
        """解析图片路径，支持本地路径和 URL。"""
        if url.startswith(('http://', 'https://')):
            try:
                resp = requests.get(url, timeout=10)
                resp.raise_for_status()
                suffix = os.path.splitext(url.split('?')[0])[1] or '.png'
                fd, path = tempfile.mkstemp(suffix=suffix)
                with os.fdopen(fd, 'wb') as f:
                    f.write(resp.content)
                return path
            except Exception as e:
                print(f'警告：图片下载失败 ({url}): {e}', file=sys.stderr)
                return None
        if os.path.isabs(url):
            return url
        return os.path.join(self.image_base_path, url)


def convert_markdown_to_docx(input_path, output_path=None):
    """将 Markdown 文件转换为 Word 文档。

    Args:
        input_path: 输入 Markdown 文件路径
        output_path: 输出 .docx 文件路径（默认与输入同名）
    Returns:
        生成的 .docx 文件路径
    """
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.docx'

    with open(input_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    document = Document()
    configure_page(document)
    configure_styles(document)

    image_base = os.path.dirname(os.path.abspath(input_path))
    converter = MarkdownToDocxConverter(document, image_base_path=image_base)
    converter.convert(md_text)

    document.save(output_path)
    return output_path
